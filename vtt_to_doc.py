from datetime import datetime, timedelta, timezone
import docx
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import webvtt
from SceneExtractor import detect_scene_changes
import cv2
from logger import logger

def timecode_to_seconds(timecode):
    """
    Converts a timecode string in the format 'HH:MM:SS' or 'MM:SS' or 'SS' to seconds.

    Args:
        timecode (str): The timecode string to convert.

    Returns:
        int: The equivalent number of seconds.

    Raises:
        ValueError: If the timecode format is invalid.
    """
    parts = timecode.split(":")
    parts = [float(part) for part in parts]
    if len(parts) == 3:
        hours, minutes, seconds = parts
    elif len(parts) == 2:
        hours = 0
        minutes, seconds = parts
    elif len(parts) == 1:
        hours = 0
        minutes = 0
        seconds = parts[0]
    else:
        raise ValueError("Invalid timecode format")
    return int(hours * 3600 + minutes * 60 + seconds)


def create_youtube_hyperlink(caption, link, format):
    """
    Creates a hyperlink to a YouTube video based on the given caption, video ID, and format.

    Args:
        caption (Caption): The caption object containing the start, end, and text of the caption.
        link (str): The link of the YouTube video.
        format (str): The format of the hyperlink. Supported formats: 'docx'.

    Returns:
        tuple: A tuple containing the hyperlink text, URL, start time, and end time.

    """

    def remove_spaces_from_text(caption):
        """
        Remove &nbsp; and similar space markers from a caption.

        Args:
            caption (str): The caption from which to remove the space markers.

        Returns:
            str: The caption with the space markers removed.
        """
        # Replace &nbsp; with a space
        caption = caption.replace("&nbsp;", " ")

        # Replace other HTML space entities
        caption = caption.replace("&ensp;", " ")
        caption = caption.replace("&emsp;", " ")
        caption = caption.replace("&thinsp;", " ")
        return caption

    start, end, text = caption.start, caption.end, caption.text
    text = remove_spaces_from_text(text)
    total_seconds = timecode_to_seconds(start)
    if format == "docx":
        # For Word, we return the URL and text separately to create a hyperlink later
        return (
            text,
            f"{link}&t={total_seconds}s",
            start,
            end,
        )


def add_hyperlink(run, url, text):
    """
    Adds a hyperlink to a run in a Word document.

    Parameters:
    - run (docx.text.run.Run): The run to add the hyperlink to.
    - url (str): The URL of the hyperlink.
    - text (str): The text to display for the hyperlink.

    Returns:
    - hyperlink (docx.oxml.shared.OxmlElement): The created hyperlink element.
    """
    # Clear the text in the original run
    run.text = ""

    r_id = run.part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    hyperlink = parse_xml(
        r'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="%s" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" />'
        % r_id
    )
    new_run = docx.oxml.shared.OxmlElement("w:r")
    run_text = docx.oxml.shared.OxmlElement("w:t")
    run_text.text = text
    new_run.append(run_text)
    hyperlink.append(new_run)
    run._r.append(hyperlink)
    return hyperlink


def should_execute_action(
    video_path, content, minutes_per_paragraph=0.5, alpha=1.0, frame_per_minute=0
):
    """
    Determines whether an action should be executed based on the time difference or the scene changes.

    Args:
        video_path (str): The path of the video file.
        content (list): A list of tuples containing the text, URL, start time, and end time.
        mode (str, optional): The mode of operation. 'time' for time difference, 'scene' for scene changes. Defaults to 'time'.
        minutes_per_paragraph (float, optional): The maximum number of minutes per paragraph. Defaults to 0.5.
        alpha (float, optional): The sensitivity of the scene change detection. Defaults to 1.0.

    Returns:
        list: A list of booleans indicating whether an action should be executed for each item in the content.
    """
    should_execute_scene = determine_execution_from_scene(
        video_path, content, alpha, frame_per_minute
    )
    should_execute_time = determine_execution_from_time(content, minutes_per_paragraph)

    should_execute = [a or b for a, b in zip(should_execute_scene, should_execute_time)]
    return should_execute


def determine_execution_from_scene(video_path, content, alpha, frame_per_minute):
    def align_timestamp(timestamp):
        return datetime.fromtimestamp(timestamp) - timedelta(hours=8)

    should_execute = []
    scene_changes = detect_scene_changes(video_path, alpha, frame_per_minute)

    logger.debug(f"scene_changes:{scene_changes}")

    i = 0
    for _, _, start, _ in content:
        current_time = parse_time(start)  # Parse the current start time
        if (current_time - align_timestamp(scene_changes[-1])) >= timedelta(minutes=0):
            should_execute.append(False)
        if i == len(scene_changes):
            should_execute.append(False)
            continue
        if (current_time - align_timestamp(scene_changes[i])) >= timedelta(minutes=0):
            i = i + 1
            should_execute.append(True)

        else:
            should_execute.append(False)
    return should_execute

def basic_execute_pattern(content):
    should_execute = []
    for _, _, _, _ in content:
        should_execute.append(False)
    return should_execute

def determine_execution_from_boolean_list(boolean_list, content):
    should_execute = []
    last_true_time = 0  # Initialize

    for i, (execute, (_, _, start, _)) in enumerate(zip(boolean_list, content)):
        # Convert the start time of the caption to seconds
        start_time = sum(
            float(x) * 60**i for i, x in enumerate(reversed(start.split(":")))
        )

        if not execute and start_time - last_true_time > 30:
            # If the current value is False and at least 30 seconds have passed since the last True,
            # set the current value to True
            should_execute.append(True)
            last_true_time = start_time  # Update the last true timestamp
        else:
            should_execute.append(execute)

    return should_execute


def determine_execution_from_time(content, minutes_per_paragraph):
    should_execute = []
    start_time = parse_time(content[0][2])  # Parse the start time
    for _, _, start, _ in content:
        current_time = parse_time(start)  # Parse the current start time
        if (current_time - start_time) >= timedelta(minutes=minutes_per_paragraph):
            should_execute.append(True)
            start_time = current_time
        else:
            should_execute.append(False)
    return should_execute


def add_frame_to_docx(cap, timestamp, para):
    # Convert timestamp to frame number
    fps = cap.get(cv2.CAP_PROP_FPS)
    timestamp_seconds = timestamp.replace(tzinfo=timezone.utc).timestamp()
    frame_number = int(fps * timestamp_seconds)
    # Set the current frame position
    cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)

    # Read the frame
    ret, frame = cap.read()
    if not ret:
        logger.debug("Failed to retrieve frame")
        return

    # Save the frame to a file
    frame_file = "frame.png"

    cv2.imwrite(frame_file, frame)

    # Add the frame to the docx file
    run = para.add_run("\n\n")
    run.add_picture(frame_file, width=Cm(15))
    return para  # Return the original Paragraph object


def write_docx(content, should_execute, picture_execute, output_file, video_path):
    """
    Executes an action (writing to a Word document) for each item in the content where should_execute is True.

    Args:
        content (list): A list of tuples containing the text, URL, start time, and end time.
        should_execute (list): A list of booleans indicating whether an action should be executed for each item in the content.
        output_file (str): The path and filename of the output Word document.

    Returns:
        None
    """
    cap = cv2.VideoCapture(video_path)
    doc = Document()
    para = doc.add_paragraph()  # Create a paragraph outside the loop
    start_time = parse_time(content[0][2])  # Parse the start time
    end_code = parse_time(content[0][3]).time()
    for (text, url, start, end), execute, pic_execute in zip(
        content, should_execute, picture_execute
    ):
        current_time = parse_time(start)
        if execute:
            para.add_run("\n")  # Insert a paragraph break
            start_code = start_time.time()
            start_code = start_code.strftime("%H:%M:%S")

            end_code = end_code.strftime("%H:%M:%S")
            para.add_run(f"({start_code} - {end_code}) \n")
            para.add_run("\n\n")  # Insert a paragraph break
            start_time = current_time
        if pic_execute:
            para = add_frame_to_docx(cap, current_time, para)
        run = para.add_run(text + " ")
        add_hyperlink(run, url, text)
        end_code = parse_time(end)

    para.add_run("\n")  # Insert a paragraph break
    para.add_run(
        f"({start_time.time().strftime('%H:%M:%S')} - {parse_time(content[-1][3]).strftime('%H:%M:%S')}) \n"
    )
    doc.save(output_file)
    cap.release()
    logger.info(f"Word file {output_file} created successfully")


def parse_time(time_str):
    """
    Parses a time string and returns a datetime object.

    Args:
        time_str (str): The time string to parse. It should be in the format "HH:MM:SS.sss" or "MM:SS.sss".

    Returns:
        datetime: A datetime object representing the parsed time.

    Raises:
        ValueError: If the time string is not in a valid format.

    """
    # Handle two types of time formats: "HH:MM:SS.sss" and "MM:SS.sss"
    if len(time_str.split(":")) == 3:
        time = datetime.strptime(time_str, "%H:%M:%S.%f")
    else:
        time = datetime.strptime(time_str, "%M:%S.%f")
    time = time.replace(year=1970)
    return time


def generate_content(vtt_file, link, format):
    """
    Generate content from a VTT file.

    Args:
        vtt_file (str): The path to the VTT file.
        link (str): The link of the video.
        format (str): The desired format of the content.

    Returns:
        list: A list of tuples containing the generated content. Each tuple consists of the text and hyperlink.

    """
    # Read the VTT file
    captions = webvtt.read(vtt_file)

    content = []
    for caption in captions:
        hyperlink = create_youtube_hyperlink(caption, link, format)
        content.append(hyperlink)  # Store the text and hyperlink as a tuple
    return content


def vtt_to_file(vtt_file, output_file, link, video_path, format,pic_embed):
    """
    Convert a VTT file to a specified format and write the content to an output file.

    Parameters:
    vtt_file (str): The path to the VTT file.
    output_file (str): The path to the output file.
    link (str): The link of the video.
    format (str): The desired format of the output file. Currently supports 'word'.

    Raises:
    ValueError: If an invalid format is provided.

    Returns:
    None
    """
    logger.debug(f"generating content")
    # Generate the content
    content = generate_content(vtt_file, link, format)

    # Write the content to the output file
    if format == "docx":
        # should_execute = should_execute_action(video_path, content, mode='scene', minutes_per_paragraph=0.5, alpha=1.0)

        if pic_embed == True:
            logger.debug(f"determine_execution_from_scene")
            should_execute_scene = determine_execution_from_scene(
                video_path, content, alpha=1, frame_per_minute=0
            )
            picture_execute = should_execute_scene
        else:
            picture_execute = basic_execute_pattern(content)
        should_execute = determine_execution_from_boolean_list(
            should_execute_scene, content
        )
        logger.debug(f"should_execute:{should_execute}")

        logger.debug(f"write_docx")
        write_docx(
            content, should_execute, picture_execute, output_file, video_path=video_path
        )
        logger.debug(f"close_vtt_to_file")

    else:
        raise ValueError("Invalid format")
