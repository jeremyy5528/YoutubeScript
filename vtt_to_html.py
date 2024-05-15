import re

import markdown
from docx import Document
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import black,blue
import webvtt
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import docx
def add_hyperlink(run, url, text):
    """
    Add a hyperlink to a run.
    """
    r_id = run.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = parse_xml(r'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="%s" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" />' % r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    run_text = docx.oxml.shared.OxmlElement('w:t')
    run_text.text = text
    new_run.append(run_text)
    hyperlink.append(new_run)
    run._r.append(hyperlink)

    return hyperlink

def write_word_file(content, output_file):
    doc = Document()
    for text, url in content:
        # Create a hyperlink in the Word document
        run = doc.add_paragraph().add_run()
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Change color to black
        run.font.underline = False  # Remove underline
        hyperlink = add_hyperlink(run, url, text)
    doc.save(output_file)
    print(f'Word file {output_file} created successfully')
def write_pdf_file(content, output_file):
    doc = SimpleDocTemplate(output_file, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    link_style = styles["BodyText"]
    link_style.textColor = blue
    link_style.underline = True
    for text, url in content:
        # Create a hyperlink in the PDF document
        link = '<link href="%s">%s</link>' % (url, text)
        story.append(Paragraph(link, link_style))
        story.append(Spacer(1, 12))
    doc.build(story)
    print(f'PDF file {output_file} created successfully')

def timecode_to_seconds(timecode):
    parts = timecode.split(':')
    parts = [float(part) for part in parts]
    if len(parts) == 3:
        hours, minutes, seconds = parts
    elif len(parts) == 2:
        hours = 0
        minutes, seconds = parts
    elif len(parts) == 1:
        hours = 0
        minutes = 0
        seconds = parts[0]
    else:
        raise ValueError("Invalid timecode format")
    return int(hours * 3600 + minutes * 60 + seconds)

def create_youtube_hyperlink(text, timecode, video_id, format):
    total_seconds = timecode_to_seconds(timecode)
    if format == 'html':
        return f'<a href="https://www.youtube.com/watch?v={video_id}&t={total_seconds}s">{text}</a>'
    elif format == 'md':
        return f'[{text}](https://www.youtube.com/watch?v={video_id}&t={total_seconds}s)'
    elif format == 'word':
        # For Word, we return the URL and text separately to create a hyperlink later
        return text, f'https://www.youtube.com/watch?v={video_id}&t={total_seconds}s'
    elif format == 'pdf':
        # For PDF, we return the URL and text separately to create a hyperlink later
        return text, f'https://www.youtube.com/watch?v={video_id}&t={total_seconds}s'

def write_html_file(html_content, output_file):
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write('\n'.join(html_content))
    print(f'HTML file {output_file} created successfully')


def write_md_file(content, output_file):
    html_content = '\n'.join(content)
    md_content = markdown.markdown(html_content)
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(md_content)
    print(f'Markdown file {output_file} created successfully')
    
def generate_content(lines, video_id, format, minutes_per_paragraph=5):
    timecode_pattern = re.compile(r'(\d{2}:\d{2}\.\d{3}) --> (\d{2}:\d{2}\.\d{3})')
    current_text = []
    start_time = None
    end_time = None
    content = ''
    paragraph_seconds = minutes_per_paragraph * 60  # Convert minutes to seconds
    current_seconds = 0
    print('test3')
    for line in lines:
        print('test2')
        line = line.strip()
        match = timecode_pattern.match(line)
        if match:
            new_seconds = timecode_to_seconds(match.group(1))
            if start_time is None:
                start_time = match.group(1)
                current_seconds = new_seconds
            elif new_seconds - current_seconds >= paragraph_seconds:
                # Add the current text and timecode to the content
                text = ' '.join(current_text)
                hyperlink = create_youtube_hyperlink(text, start_time, video_id, format)
                if format == 'html':
                    content += f'<p>{start_time} - {end_time}<br>{hyperlink}</p>'
                elif format == 'md':
                    content += f'\n**{start_time} - {end_time}**\n{hyperlink}\n'
                elif format == 'word':
                    content += f'{start_time} - {end_time}\n{hyperlink}\n'
                current_text = []
                start_time = match.group(1)
                current_seconds = new_seconds
            end_time = match.group(2)
        elif line:
            current_text.append(line)
    
    if current_text and start_time:
        text = ' '.join(current_text)
        print('test')
        hyperlink = create_youtube_hyperlink(text, start_time, video_id, format)
        if format == 'html':
            content += f'<p>{start_time} - {end_time}<br>{hyperlink}</p>'
        elif format == 'md':
            content += f'\n**{start_time} - {end_time}**\n{hyperlink}\n'
        elif format == 'word':
            content += f'{start_time} - {end_time}\n{hyperlink}\n'
    return content

def vtt_to_file(vtt_file, output_file, video_id, format):
    # Read the VTT file
    captions = webvtt.read(vtt_file)

    # Generate the content
    content = []
    for caption in captions:
        start = caption.start
        text = caption.text
        hyperlink = create_youtube_hyperlink(text, start, video_id, format)
        content.append(hyperlink)

    # Write the content to the output file
    if format == 'html':
        write_html_file(content, output_file)
    elif format == 'pdf':
        write_pdf_file(content, output_file)
    elif format == 'word':
        write_word_file(content, output_file)
    elif format == 'md':
        write_md_file(content, output_file)
    else:
        raise ValueError("Invalid format")

# Example usage
vtt_file = "D:\\youtubescript\\text\\【人工智能】Google发布AlphaFold 3 ｜ 可预测所有生命分子组合 ｜ 准确性提高100% ｜ AlphaFold Server ｜ DeepMind ｜ Isomorphic Labs.vtt"
output_file = 'output'
video_id = 'MszokltHN0Q'
vtt_to_file(vtt_file, output_file + '.html', video_id, 'html')
vtt_to_file(vtt_file, output_file + '.pdf', video_id, 'pdf')
vtt_to_file(vtt_file, output_file + '.docx', video_id, 'word')
vtt_to_file(vtt_file, output_file + '.md', video_id, 'md')